"""
Local file-system tools (run on the server machine).

Useful for desktop workflows like "save this to a file", "what's in this
folder", "open this file". Delete is marked dangerous so it needs confirmation.
"""

from __future__ import annotations

import logging
import os
import subprocess
import time

from app.services.agent.tool_registry import tool

logger = logging.getLogger("J.A.R.V.I.S")

_MAX_READ = 20000  # chars

# Common user folders that a person refers to by bare name ("desktop pe ...").
# Mapping a leading 'desktop'/'downloads'/etc. to the real home folder is NOT a
# hardcoded app behaviour -- it's just resolving the well-known Windows user
# directories so natural commands work.
_USER_DIRS = {
    "desktop", "downloads", "documents", "pictures", "music", "videos", "desktop",
}


def _resolve(path: str) -> str:
    """Resolve a user-given path: expand ~ and env vars, and treat a leading
    well-known folder name (desktop/downloads/...) as relative to the home dir.
    """
    p = (path or "").strip().strip('"').strip("'")
    if not p:
        return ""
    p = os.path.expandvars(os.path.expanduser(p))
    if not os.path.isabs(p):
        norm = p.replace("\\", "/")
        first = norm.split("/")[0].lower()
        if first in _USER_DIRS:
            home = os.path.expanduser("~")
            return os.path.normpath(os.path.join(home, *norm.split("/")))
    return p


# --------------------------------------------------------------------------- #
# accident guard
# --------------------------------------------------------------------------- #
# NOT a security boundary -- this is a single-user local app and the user can
# delete anything they like through Explorer. It exists because the *model*
# chooses these paths: one hallucinated argument used to be enough to delete
# from C:\Windows or overwrite part of JARVIS itself. Reads are never blocked;
# only writes and deletes.
def _protected_roots() -> list:
    roots = []
    for var in ("SystemRoot", "ProgramFiles", "ProgramFiles(x86)", "ProgramData"):
        value = os.environ.get(var)
        if value:
            roots.append(os.path.normcase(os.path.normpath(value)))
    # JARVIS' own source tree: an agent editing its own running code mid-turn
    # is never what the user meant.
    try:
        from config import BASE_DIR
        roots.append(os.path.normcase(os.path.normpath(str(BASE_DIR))))
    except Exception:  # noqa: BLE001
        pass
    return roots


def _guard_path(path: str, action: str = "modify") -> str:
    """Return an ERROR string when a write/delete target looks like a mistake.

    Returns "" when the path is fine, so callers read as:
        blocked = _guard_path(p, "delete")
        if blocked:
            return blocked
    """
    if not path:
        return ""
    try:
        target = os.path.normcase(os.path.abspath(path))
    except (OSError, ValueError):
        return ""

    # A drive root ("C:\\") is never a deliberate target.
    drive, tail = os.path.splitdrive(target)
    if drive and tail.strip("\\/") == "":
        return f"ERROR: refusing to {action} a drive root ({path})."

    for root in _protected_roots():
        if target == root or target.startswith(root + os.sep):
            return (f"ERROR: refusing to {action} '{path}' -- it is inside a "
                    f"protected system or application folder. If you really "
                    f"meant this, do it manually in File Explorer.")
    return ""


@tool(
    name="list_directory",
    description="List the files and folders inside a directory on the computer.",
    params={"path": {"type": "string", "description": "Directory path to list."}},
    category="desktop",
    verification={"family": "query", "cacheable": False},
)
def list_directory(path: str) -> str:
    p = _resolve(path)
    if not p:
        return "ERROR: no path given."
    if not os.path.isdir(p):
        return f"ERROR: '{p}' is not a directory."
    try:
        entries = sorted(os.listdir(p))
        if not entries:
            return f"'{p}' is empty."
        lines = []
        for name in entries[:200]:
            full = os.path.join(p, name)
            kind = "DIR " if os.path.isdir(full) else "FILE"
            lines.append(f"[{kind}] {name}")
        more = "" if len(entries) <= 200 else f"\n... and {len(entries) - 200} more"
        return f"Contents of {p}:\n" + "\n".join(lines) + more
    except Exception as e:  # noqa: BLE001
        return f"ERROR: could not list '{p}': {e}"


@tool(
    name="read_file",
    description="Read and return the text content of a file on the computer.",
    params={"path": {"type": "string", "description": "File path to read."}},
    category="desktop",
    verification={"family": "query", "cacheable": False},
)
def read_file(path: str) -> str:
    p = _resolve(path)
    if not os.path.isfile(p):
        return f"ERROR: '{p}' is not a file."
    try:
        with open(p, "r", encoding="utf-8", errors="replace") as f:
            data = f.read(_MAX_READ + 1)
        truncated = len(data) > _MAX_READ
        return data[:_MAX_READ] + ("\n... [truncated]" if truncated else "")
    except Exception as e:  # noqa: BLE001
        return f"ERROR: could not read '{p}': {e}"


@tool(
    name="write_file",
    description=(
        "Write text content to a file on the computer (creates or overwrites). "
        "Use for saving notes, code, or generated text to disk."
    ),
    params={
        "path": {"type": "string", "description": "File path to write to."},
        "content": {"type": "string", "description": "Text content to write."},
    },
    category="desktop",
    verification={"family": "file"},
)
def write_file(path: str, content: str) -> str:
    p = _resolve(path)
    if not p:
        return "ERROR: no path given."
    blocked = _guard_path(p, "write to")
    if blocked:
        return blocked
    try:
        parent = os.path.dirname(p)
        if parent:
            os.makedirs(parent, exist_ok=True)
        with open(p, "w", encoding="utf-8") as f:
            f.write(content or "")
        logger.info("[FILE] Wrote %d chars to %s", len(content or ""), p)
        return f"Saved {len(content or '')} characters to {p}."
    except Exception as e:  # noqa: BLE001
        return f"ERROR: could not write '{p}': {e}"


@tool(
    name="open_file",
    description="Open a file or folder with its default application on Windows.",
    params={"path": {"type": "string", "description": "File or folder path to open."}},
    category="desktop",
    verification={"family": "open"},
)
def open_file(path: str) -> str:
    p = _resolve(path)
    if not os.path.exists(p):
        return f"ERROR: '{p}' does not exist."
    try:
        os.startfile(p)  # type: ignore[attr-defined]  # Windows-only
        return f"Opened {p}."
    except Exception as e:  # noqa: BLE001
        return f"ERROR: could not open '{p}': {e}"


@tool(
    name="delete_file",
    description=(
        "Delete a file from the computer. Destructive and permanent. "
        "Only call after the user has confirmed."
    ),
    params={"path": {"type": "string", "description": "File path to delete."}},
    dangerous=True,
    category="desktop",
    verification={"family": "file"},
)
def delete_file(path: str) -> str:
    p = _resolve(path)
    if not os.path.isfile(p):
        return f"ERROR: '{p}' is not a file."
    blocked = _guard_path(p, "delete")
    if blocked:
        return blocked
    try:
        os.remove(p)
        logger.info("[FILE] Deleted %s", p)
        return f"Deleted {p}."
    except Exception as e:  # noqa: BLE001
        return f"ERROR: could not delete '{p}': {e}"


@tool(
    name="create_folder",
    description=(
        "Create a new folder / directory on the computer. Handles common roots "
        "like Desktop, Downloads, Documents by name. Example: path "
        "'desktop/ayushman' makes a folder named 'ayushman' on the Desktop."
    ),
    params={"path": {"type": "string", "description": "Folder path to create, e.g. 'desktop/ayushman'."}},
    category="desktop",
    verification={"family": "file"},
)
def create_folder(path: str) -> str:
    p = _resolve(path)
    if not p:
        return "ERROR: no folder path given."
    blocked = _guard_path(p, "create a folder in")
    if blocked:
        return blocked
    try:
        if os.path.isdir(p):
            return f"Folder already exists: {p}."
        os.makedirs(p, exist_ok=True)
        logger.info("[FILE] Created folder %s", p)
        return f"Created folder {p}."
    except Exception as e:  # noqa: BLE001
        return f"ERROR: could not create folder '{p}': {e}"


@tool(
    name="move_to_trash",
    description=(
        "Move a file or folder to the Recycle Bin (recoverable, safer than "
        "delete). Use for 'recycle bin me daalo' / 'trash me bhejo'."
    ),
    params={"path": {"type": "string", "description": "File or folder path to move to the Recycle Bin."}},
    category="desktop",
    verification={"family": "file"},
)
def move_to_trash(path: str) -> str:
    p = _resolve(path)
    if not p:
        return "ERROR: no path given."
    if not os.path.exists(p):
        return f"ERROR: '{p}' does not exist."
    blocked = _guard_path(p, "recycle")
    if blocked:
        return blocked
    try:
        from send2trash import send2trash
        send2trash(p)
        logger.info("[FILE] Sent to Recycle Bin: %s", p)
        return f"Moved {p} to the Recycle Bin."
    except ImportError:
        return (
            "ERROR: Recycle Bin support needs the 'send2trash' package "
            "(pip install send2trash)."
        )
    except Exception as e:  # noqa: BLE001
        return f"ERROR: could not move '{p}' to the Recycle Bin: {e}"


@tool(
    name="move_path",
    description=(
        "Move or rename a file or folder from a source path to a destination "
        "path. Handles Desktop/Downloads/Documents by name."
    ),
    params={
        "source": {"type": "string", "description": "Existing file/folder path."},
        "destination": {"type": "string", "description": "New path or folder to move into."},
    },
    category="desktop",
    verification={"family": "file"},
)
def move_path(source: str, destination: str) -> str:
    src = _resolve(source)
    dst = _resolve(destination)
    if not src or not dst:
        return "ERROR: both source and destination are required."
    if not os.path.exists(src):
        return f"ERROR: '{src}' does not exist."
    for candidate, verb in ((src, "move"), (dst, "move into")):
        blocked = _guard_path(candidate, verb)
        if blocked:
            return blocked
    try:
        import shutil
        # If destination is an existing directory, move INTO it.
        if os.path.isdir(dst):
            dst = os.path.join(dst, os.path.basename(src.rstrip("/\\")))
        parent = os.path.dirname(dst)
        if parent:
            os.makedirs(parent, exist_ok=True)
        shutil.move(src, dst)
        logger.info("[FILE] Moved %s -> %s", src, dst)
        return f"Moved to {dst}."
    except Exception as e:  # noqa: BLE001
        return f"ERROR: could not move '{src}': {e}"


@tool(
    name="find_file",
    description=(
        "Find a file anywhere in the user's own folders (Desktop, Documents, "
        "Downloads, Pictures, Videos, Music) by part of its name. Use this when "
        "the user asks where something is, or refers to a file without giving a "
        "path -- 'that resume PDF', 'the budget sheet'. Returns full paths that "
        "can be passed straight to read_file, read_document or open_file."
    ),
    params={
        "name": {"type": "string",
                 "description": "Part of the file name to search for, e.g. 'resume'."},
        "extension": {"type": "string", "required": False,
                      "description": "Optional file type filter, e.g. 'pdf' or '.docx'."},
    },
    category="desktop",
    verification={"family": "query", "cacheable": False},
)
def find_file(name: str, extension: str = "") -> str:
    query = (name or "").strip()
    if not query:
        return "ERROR: no name to search for."
    try:
        from app.services.agent.file_index import get_file_index
        index = get_file_index()
    except Exception as e:  # noqa: BLE001
        return f"ERROR: file search is unavailable: {e}"

    if not index.enabled:
        return "ERROR: the file index could not be opened."

    hits = index.search(query, limit=20, extension=extension)
    if not hits:
        if index.count() == 0:
            index.build_async()
            return "The file index is still being built -- ask again in a few seconds."
        where = ", ".join(os.path.basename(r) for r in index.roots())
        return f"No file matching '{query}' in {where}."

    lines = []
    for hit in hits:
        size_kb = (hit["size"] or 0) / 1024.0
        stamp = time.strftime("%Y-%m-%d", time.localtime(hit["mtime"] or 0))
        lines.append(f"{hit['path']}  ({size_kb:,.0f} KB, {stamp})")
    return f"{len(hits)} match(es) for '{query}':\n" + "\n".join(lines)


@tool(
    name="read_document",
    description=(
        "Read the text out of a PDF, Word (.docx) or Excel (.xlsx) file. "
        "read_file only handles plain text -- use this one for documents."
    ),
    params={
        "path": {"type": "string", "description": "Path to the .pdf, .docx or .xlsx file."},
        "max_pages": {"type": "int", "required": False,
                      "description": "For PDFs: how many pages to read (default 20)."},
    },
    category="desktop",
    verification={"family": "query", "cacheable": False},
)
def read_document(path: str, max_pages: int = 20) -> str:
    p = _resolve(path)
    if not os.path.isfile(p):
        return f"ERROR: '{p}' is not a file."
    ext = os.path.splitext(p)[1].lower()
    try:
        if ext == ".pdf":
            text = _read_pdf(p, max_pages)
        elif ext == ".docx":
            text = _read_docx(p)
        elif ext in (".xlsx", ".xlsm"):
            text = _read_xlsx(p)
        elif ext == ".doc":
            return ("ERROR: legacy .doc is not supported. Open it in Word and "
                    "save it as .docx first.")
        else:
            return (f"ERROR: read_document handles .pdf, .docx and .xlsx. "
                    f"For '{ext or 'this type'}' use read_file instead.")
    except ImportError as e:
        return f"ERROR: {e}"
    except Exception as e:  # noqa: BLE001
        return f"ERROR: could not read '{p}': {e}"

    if not text.strip():
        return (f"'{os.path.basename(p)}' has no extractable text. It may be a "
                f"scan -- open it and try read_screen instead.")
    truncated = len(text) > _MAX_READ
    return text[:_MAX_READ] + ("\n... [truncated]" if truncated else "")


def _read_pdf(path: str, max_pages: int) -> str:
    try:
        from pypdf import PdfReader
    except ImportError:
        try:
            from PyPDF2 import PdfReader  # type: ignore[no-redef]
        except ImportError:
            raise ImportError("reading PDFs needs the 'pypdf' package "
                              "(pip install pypdf).") from None
    reader = PdfReader(path)
    limit = max(1, int(max_pages or 20))
    out = []
    for number, page in enumerate(reader.pages[:limit], 1):
        try:
            body = page.extract_text() or ""
        except Exception:  # noqa: BLE001 - one bad page must not lose the rest
            body = ""
        if body.strip():
            out.append(f"--- page {number} ---\n{body.strip()}")
    if len(reader.pages) > limit:
        out.append(f"... {len(reader.pages) - limit} more page(s) not read.")
    return "\n\n".join(out)


def _read_docx(path: str) -> str:
    try:
        import docx
    except ImportError:
        raise ImportError("reading Word files needs the 'python-docx' package "
                          "(pip install python-docx).") from None
    document = docx.Document(path)
    parts = [para.text for para in document.paragraphs if para.text.strip()]
    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            if any(cells):
                parts.append(" | ".join(cells))
    return "\n".join(parts)


def _read_xlsx(path: str) -> str:
    try:
        import openpyxl
    except ImportError:
        raise ImportError("reading Excel files needs the 'openpyxl' package "
                          "(pip install openpyxl).") from None
    workbook = openpyxl.load_workbook(path, read_only=True, data_only=True)
    out = []
    for sheet in workbook.worksheets:
        out.append(f"--- sheet: {sheet.title} ---")
        for row in sheet.iter_rows(values_only=True):
            cells = [str(v) for v in row if v is not None]
            if cells:
                out.append(" | ".join(cells))
    workbook.close()
    return "\n".join(out)


@tool(
    name="zip_files",
    description=(
        "Compress a file or a whole folder into a .zip archive. Use for "
        "'zip this folder' / 'compress these files'."
    ),
    params={
        "source": {"type": "string", "description": "File or folder to compress."},
        "destination": {"type": "string", "required": False,
                        "description": "Output .zip path. Defaults to next to the source."},
    },
    category="desktop",
    verification={"family": "file"},
)
def zip_files(source: str, destination: str = "") -> str:
    src = _resolve(source)
    if not src or not os.path.exists(src):
        return f"ERROR: '{source}' does not exist."
    dst = _resolve(destination) if destination else src.rstrip("/\\") + ".zip"
    if not dst.lower().endswith(".zip"):
        dst += ".zip"
    blocked = _guard_path(dst, "write to")
    if blocked:
        return blocked
    try:
        import zipfile
        parent = os.path.dirname(dst)
        if parent:
            os.makedirs(parent, exist_ok=True)
        count = 0
        with zipfile.ZipFile(dst, "w", zipfile.ZIP_DEFLATED) as archive:
            if os.path.isfile(src):
                archive.write(src, os.path.basename(src))
                count = 1
            else:
                base = os.path.dirname(src.rstrip("/\\"))
                for dirpath, _dirnames, filenames in os.walk(src):
                    for filename in filenames:
                        full = os.path.join(dirpath, filename)
                        archive.write(full, os.path.relpath(full, base))
                        count += 1
        size_kb = os.path.getsize(dst) / 1024.0
        logger.info("[FILE] Zipped %d file(s) -> %s", count, dst)
        return f"Compressed {count} file(s) into {dst} ({size_kb:,.0f} KB)."
    except Exception as e:  # noqa: BLE001
        return f"ERROR: could not create '{dst}': {e}"


@tool(
    name="unzip_file",
    description="Extract a .zip archive into a folder.",
    params={
        "path": {"type": "string", "description": "The .zip file to extract."},
        "destination": {"type": "string", "required": False,
                        "description": "Folder to extract into. Defaults to a folder beside the zip."},
    },
    category="desktop",
    verification={"family": "file"},
)
def unzip_file(path: str, destination: str = "") -> str:
    src = _resolve(path)
    if not os.path.isfile(src):
        return f"ERROR: '{path}' is not a file."
    dst = _resolve(destination) if destination else os.path.splitext(src)[0]
    blocked = _guard_path(dst, "extract into")
    if blocked:
        return blocked
    try:
        import zipfile
        os.makedirs(dst, exist_ok=True)
        root = os.path.abspath(dst)
        with zipfile.ZipFile(src) as archive:
            # Reject entries that would escape the destination (zip-slip).
            for member in archive.namelist():
                target = os.path.abspath(os.path.join(dst, member))
                if target != root and not target.startswith(root + os.sep):
                    return (f"ERROR: '{os.path.basename(src)}' contains an entry "
                            f"that would write outside the destination ({member}).")
            archive.extractall(dst)
            count = len(archive.namelist())
        logger.info("[FILE] Extracted %d entry(ies) -> %s", count, dst)
        return f"Extracted {count} item(s) into {dst}."
    except Exception as e:  # noqa: BLE001
        return f"ERROR: could not extract '{src}': {e}"
